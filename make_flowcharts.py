import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import os

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── colours ──────────────────────────────────────────────────────────────────
C_TRIGGER  = '#1F497D'   # dark blue  — trigger box
C_DECISION = '#FFC000'   # amber      — decision diamond
C_ACTION   = '#4472C4'   # mid blue   — action / state box
C_TERM     = '#70AD47'   # green      — terminal (idle / done)
C_ABORT    = '#FF0000'   # red        — abort / error
C_MANUAL   = '#7030A0'   # purple     — manual mode
C_TEXT_LT  = '#FFFFFF'
C_TEXT_DK  = '#1F1F1F'
C_ARROW    = '#404040'

# ── drawing helpers ───────────────────────────────────────────────────────────

def box(ax, cx, cy, w, h, text, color, textcolor=C_TEXT_LT, fontsize=8.5, bold=False):
    ax.add_patch(FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                                boxstyle='round,pad=0.02',
                                facecolor=color, edgecolor='white', linewidth=1.2, zorder=3))
    ax.text(cx, cy, text, ha='center', va='center', fontsize=fontsize,
            color=textcolor, fontweight='bold' if bold else 'normal',
            wrap=True, zorder=4, multialignment='center')

def diamond(ax, cx, cy, w, h, text, color=C_DECISION, textcolor=C_TEXT_DK, fontsize=8):
    xs = [cx, cx + w/2, cx, cx - w/2, cx]
    ys = [cy + h/2, cy, cy - h/2, cy, cy + h/2]
    ax.fill(xs, ys, color=color, edgecolor='white', linewidth=1.2, zorder=3)
    ax.text(cx, cy, text, ha='center', va='center', fontsize=fontsize,
            color=textcolor, fontweight='bold', zorder=4, multialignment='center')

def arrow(ax, x1, y1, x2, y2, label='', color=C_ARROW, lw=1.5):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw), zorder=2)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx + 0.05, my, label, fontsize=7.5, color=color, va='center')

def save_chart(fig, name):
    path = os.path.join(os.path.dirname(__file__), name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='#F8F8F8')
    plt.close(fig)
    return path

# ─────────────────────────────────────────────────────────────────────────────
# 1.  UPSHIFT
# ─────────────────────────────────────────────────────────────────────────────
def chart_upshift():
    fig, ax = plt.subplots(figsize=(6, 10))
    ax.set_xlim(0, 6); ax.set_ylim(0, 10); ax.axis('off')
    ax.set_facecolor('#F8F8F8')
    fig.patch.set_facecolor('#F8F8F8')
    ax.set_title('UPSHIFT BUTTON PRESSED', fontsize=13, fontweight='bold', color=C_TRIGGER, pad=10)

    # nodes  (cx, cy, w, h)
    box(ax,    3, 9.3, 3.2, 0.5, 'SHIFT UP button pressed', C_TRIGGER, bold=True)
    diamond(ax,3, 8.4, 3.5, 0.7, 'Shift in progress?')
    box(ax,    3, 7.4, 3.2, 0.5, 'IGNORED — dropped', C_ABORT)

    diamond(ax,3, 6.5, 3.5, 0.7, 'Current gear?')

    box(ax,    3, 5.5, 3.2, 0.5, 'IGNORED (top gear)', C_ABORT)

    box(ax,    3, 4.5, 3.2, 0.55,'State: UPSHIFTING\nSend IGN cut + CAN upshift', C_ACTION)
    box(ax,    3, 3.5, 3.2, 0.55,'Wait relay timer\n(shiftUpMs)', C_ACTION)
    diamond(ax,3, 2.5, 3.5, 0.7, 'CAN gear pos\narrived?')
    box(ax,    3, 1.5, 3.2, 0.55,'State: IDLE_GEAR_x\n(corrected by CAN)', C_TERM)

    arrow(ax, 3, 9.05, 3, 8.75)
    arrow(ax, 3, 8.05, 3, 7.65,  'YES')
    # NO branch — go to gear check
    ax.annotate('', xy=(3, 6.85), xytext=(3, 8.05),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(3.12, 7.45, 'NO', fontsize=7.5, color=C_ARROW)

    arrow(ax, 3, 6.15, 3, 5.75,  'GEAR 6')
    # Gears 1-5 branch
    ax.annotate('', xy=(3, 4.78), xytext=(3, 6.15),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(3.12, 5.45, 'GEAR 1-5', fontsize=7.5, color=C_ARROW)

    arrow(ax, 3, 4.22, 3, 3.78)
    arrow(ax, 3, 3.22, 3, 2.85)
    arrow(ax, 3, 2.15, 3, 1.78,  'YES')
    # NO loop
    ax.annotate('', xy=(1.8, 2.5), xytext=(3, 2.15),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(1.8, 3.5), xytext=(1.8, 2.5),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(3, 3.78), xytext=(1.8, 3.5),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(1.35, 3.0, 'NO\n(wait)', fontsize=7, color=C_ARROW, ha='center')

    return save_chart(fig, '_fc_upshift.png')


# ─────────────────────────────────────────────────────────────────────────────
# 2.  DOWNSHIFT
# ─────────────────────────────────────────────────────────────────────────────
def chart_downshift():
    fig, ax = plt.subplots(figsize=(7, 13))
    ax.set_xlim(0, 7); ax.set_ylim(0, 13); ax.axis('off')
    ax.set_facecolor('#F8F8F8'); fig.patch.set_facecolor('#F8F8F8')
    ax.set_title('DOWNSHIFT BUTTON PRESSED', fontsize=13, fontweight='bold', color=C_TRIGGER, pad=10)

    box(ax,    3.5, 12.3, 3.5, 0.5, 'SHIFT DOWN button pressed', C_TRIGGER, bold=True)
    diamond(ax,3.5, 11.4, 3.8, 0.7, 'Shift in progress?')
    box(ax,    3.5, 10.4, 3.2, 0.5, 'IGNORED — dropped', C_ABORT)

    diamond(ax,3.5,  9.4, 3.8, 0.7, 'Current gear?')
    box(ax,    3.5,  8.4, 3.2, 0.5, 'IGNORED (already in 1st)', C_ABORT)

    # --- Gear 1 path (ignored) already above
    # --- Neutral path (right side)
    box(ax,    6.0,  7.5, 1.8, 0.55,'State:\nWAITING_FOR\n_CLUTCH\n_SHIFT_DOWN', C_ACTION, fontsize=7)
    diamond(ax,6.0,  6.1, 1.8, 0.7, 'Clutch\npulled?', fontsize=7.5)
    box(ax,    6.0,  4.9, 1.8, 0.55,'NEUTRAL_DOWN\n_SHIFTING\nCAN downshift\n(no IGN cut)', C_ACTION, fontsize=7)
    box(ax,    6.0,  3.8, 1.8, 0.5, 'IDLE_NEUTRAL\n(abort)', C_ABORT, fontsize=7.5)

    # --- Gear 2-6 main path
    box(ax,    3.0,  7.5, 3.0, 0.55,'State:\nDOWNSHIFT_CLUTCH_ENGAGING\nEngage clutch servo', C_ACTION, fontsize=7.5)
    diamond(ax,3.0,  6.1, 3.0, 0.7, 'Clutch pulled\nwithin 200 ms?')
    box(ax,    3.0,  4.9, 3.0, 0.55,'State: DOWNSHIFT_SHIFTING\nSend CAN downshift cmd', C_ACTION)
    box(ax,    3.0,  3.8, 3.0, 0.5, 'Wait relay timer (shiftDownMs)', C_ACTION)
    diamond(ax,3.0,  2.8, 3.0, 0.7, 'CAN gear pos\narrived?')
    box(ax,    3.0,  1.8, 3.0, 0.55,'State: IDLE_GEAR_x\n(corrected by CAN)', C_TERM)

    # arrows — main column
    arrow(ax, 3.5, 12.05, 3.5, 11.75)
    arrow(ax, 3.5, 11.05, 3.5, 10.65, 'YES')
    ax.annotate('', xy=(3.5, 9.75), xytext=(3.5, 11.05),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(3.62, 10.35, 'NO', fontsize=7.5, color=C_ARROW)

    arrow(ax, 3.5, 9.05, 3.5, 8.65, 'GEAR 1')
    # Neutral branch right
    ax.annotate('', xy=(6.0, 7.78), xytext=(3.5, 9.05),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(5.1, 8.55, 'NEUTRAL', fontsize=7.5, color=C_ARROW)
    # Gear 2-6 branch
    ax.annotate('', xy=(3.0, 7.78), xytext=(3.5, 9.05),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(2.3, 8.55, 'GEAR 2-6', fontsize=7.5, color=C_ARROW)

    # Neutral sub-flow
    arrow(ax, 6.0, 7.22, 6.0, 6.45)
    arrow(ax, 6.0, 5.75, 6.0, 5.18, 'YES')
    ax.annotate('', xy=(6.0, 4.05), xytext=(6.0, 5.75),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(6.12, 4.8, 'NO\n(200ms)', fontsize=7, color=C_ARROW)

    # Gear 2-6 sub-flow
    arrow(ax, 3.0, 7.22, 3.0, 6.45)
    ax.annotate('', xy=(3.0, 5.18), xytext=(3.0, 5.75),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(3.12, 5.45, 'YES', fontsize=7.5, color=C_ARROW)
    # timeout fallback arrow
    ax.annotate('', xy=(1.5, 6.1), xytext=(3.0, 5.75),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(1.5, 4.9), xytext=(1.5, 6.1),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(3.0, 5.18), xytext=(1.5, 4.9),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(0.7, 5.5, 'NO\n(200ms\nfallback)', fontsize=7, color=C_ARROW, ha='center')

    arrow(ax, 3.0, 4.62, 3.0, 4.05)
    arrow(ax, 3.0, 3.55, 3.0, 3.15)
    arrow(ax, 3.0, 2.45, 3.0, 2.08, 'YES')
    ax.annotate('', xy=(1.2, 2.8), xytext=(3.0, 2.45),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(1.2, 3.8), xytext=(1.2, 2.8),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.annotate('', xy=(3.0, 4.05), xytext=(1.2, 3.8),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(0.6, 3.3, 'NO\n(wait)', fontsize=7, color=C_ARROW, ha='center')

    return save_chart(fig, '_fc_downshift.png')


# ─────────────────────────────────────────────────────────────────────────────
# 3.  NEUTRAL BUTTON
# ─────────────────────────────────────────────────────────────────────────────
def chart_neutral():
    fig, ax = plt.subplots(figsize=(7, 11))
    ax.set_xlim(0, 7); ax.set_ylim(0, 11); ax.axis('off')
    ax.set_facecolor('#F8F8F8'); fig.patch.set_facecolor('#F8F8F8')
    ax.set_title('NEUTRAL BUTTON PRESSED', fontsize=13, fontweight='bold', color=C_TRIGGER, pad=10)

    box(ax,    3.5, 10.3, 3.5, 0.5,  'NEUTRAL button pressed\n(single button, auto direction)', C_TRIGGER, bold=True)
    diamond(ax,3.5,  9.3, 3.8, 0.7,  'Shift in progress?')
    box(ax,    3.5,  8.3, 3.2, 0.5,  'IGNORED — dropped', C_ABORT)

    diamond(ax,3.5,  7.3, 3.8, 0.7,  'Current gear?')
    box(ax,    3.5,  6.3, 3.2, 0.5,  'IGNORED\n(not gear 1 or 2)', C_ABORT)

    # LEFT: gear 1 path → N
    box(ax,    1.6,  5.2, 2.4, 0.55, 'State:\nWAITING_FOR\n_CLUTCH_NEUTRAL_UP\n(1st → N)', C_ACTION, fontsize=7.5)
    diamond(ax,1.6,  3.9, 2.4, 0.7,  'Clutch\npulled?', fontsize=8)
    box(ax,    1.6,  2.8, 2.4, 0.55, 'NEUTRAL_UP_SHIFTING\nCAN upshift\n(no IGN cut)', C_ACTION, fontsize=7.5)
    box(ax,    1.6,  1.7, 2.4, 0.5,  'IDLE_NEUTRAL\n(corrected by CAN)', C_TERM, fontsize=7.5)
    box(ax,    0.2,  2.8, 0.8, 0.5,  'ABORT\n(IDLE_NEUTRAL)', C_ABORT, fontsize=7)

    # RIGHT: gear 2 path → N
    box(ax,    5.4,  5.2, 2.4, 0.55, 'State:\nWAITING_FOR\n_CLUTCH_NEUTRAL_DOWN\n(2nd → N)', C_ACTION, fontsize=7.5)
    diamond(ax,5.4,  3.9, 2.4, 0.7,  'Clutch\npulled?', fontsize=8)
    box(ax,    5.4,  2.8, 2.4, 0.55, 'NEUTRAL_DOWN_SHIFTING\nCAN downshift\n(no IGN cut)', C_ACTION, fontsize=7.5)
    box(ax,    5.4,  1.7, 2.4, 0.5,  'IDLE_NEUTRAL\n(corrected by CAN)', C_TERM, fontsize=7.5)
    box(ax,    6.8,  2.8, 0.8, 0.5,  'ABORT\n(IDLE_NEUTRAL)', C_ABORT, fontsize=7)

    # main arrows
    arrow(ax, 3.5, 10.05, 3.5, 9.65)
    arrow(ax, 3.5, 8.95,  3.5, 8.55, 'YES')
    ax.annotate('', xy=(3.5, 7.65), xytext=(3.5, 8.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(3.62, 8.2, 'NO', fontsize=7.5, color=C_ARROW)

    arrow(ax, 3.5, 6.95,  3.5, 6.55, 'OTHER')
    ax.annotate('', xy=(1.6, 5.48), xytext=(3.5, 6.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(2.1, 6.4, 'GEAR 1', fontsize=7.5, color=C_ARROW)
    ax.annotate('', xy=(5.4, 5.48), xytext=(3.5, 6.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(4.6, 6.4, 'GEAR 2', fontsize=7.5, color=C_ARROW)

    # left sub-flow
    arrow(ax, 1.6, 4.95,  1.6, 4.25)
    arrow(ax, 1.6, 3.55,  1.6, 3.08, 'YES')
    ax.annotate('', xy=(0.2, 3.9), xytext=(1.6, 3.55),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(0.65, 3.62, 'NO\n200ms', fontsize=7, color=C_ARROW)
    arrow(ax, 1.6, 2.52,  1.6, 1.95)

    # right sub-flow
    arrow(ax, 5.4, 4.95,  5.4, 4.25)
    arrow(ax, 5.4, 3.55,  5.4, 3.08, 'YES')
    ax.annotate('', xy=(6.8, 3.9), xytext=(5.4, 3.55),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(6.0, 3.62, 'NO\n200ms', fontsize=7, color=C_ARROW)
    arrow(ax, 5.4, 2.52,  5.4, 1.95)

    ax.text(1.6, 0.8, 'Relay done → IDLE corrected\nby CAN gear position', fontsize=7,
            ha='center', color='#404040', style='italic')
    ax.text(5.4, 0.8, 'Relay done → IDLE corrected\nby CAN gear position', fontsize=7,
            ha='center', color='#404040', style='italic')

    return save_chart(fig, '_fc_neutral.png')


# ─────────────────────────────────────────────────────────────────────────────
# 4.  MANUAL MODE
# ─────────────────────────────────────────────────────────────────────────────
def chart_manual():
    fig, ax = plt.subplots(figsize=(8, 11))
    ax.set_xlim(0, 8); ax.set_ylim(0, 11); ax.axis('off')
    ax.set_facecolor('#F8F8F8'); fig.patch.set_facecolor('#F8F8F8')
    ax.set_title('MANUAL MODE', fontsize=13, fontweight='bold', color=C_MANUAL, pad=10)

    # Toggle section
    box(ax,    4.0, 10.3, 4.0, 0.5,  'MANUAL TOGGLE button held (GPIO 10)', C_MANUAL, bold=True)
    diamond(ax,4.0,  9.3, 4.0, 0.7,  'Held for 1000 ms?')
    box(ax,    4.0,  8.3, 3.5, 0.5,  'IGNORED — released too soon', C_ABORT)
    diamond(ax,4.0,  7.3, 4.0, 0.7,  'Currently in manual mode?')
    box(ax,    1.8,  6.3, 2.8, 0.5,  'ENTER MANUAL MODE\nAmber flash on matrix', C_MANUAL)
    box(ax,    6.2,  6.3, 2.8, 0.5,  'EXIT MANUAL MODE\nReturn to normal', C_TERM)

    # Separator
    ax.axhline(5.6, color='#CCCCCC', lw=1, linestyle='--')
    ax.text(4.0, 5.4, 'While in Manual Mode — state machine BYPASSED — direct CAN commands',
            ha='center', fontsize=8, color='#606060', style='italic')

    # Manual shift up
    box(ax,    1.8, 4.7, 2.8, 0.5,  'SHIFT UP button pressed', C_MANUAL, bold=True)
    box(ax,    1.8, 3.8, 2.8, 0.55, 'Send CAN upshift\n100 ms pulse + IGN cut\ntargetGear = UNKNOWN', C_ACTION)
    box(ax,    1.8, 2.8, 2.8, 0.5,  'Relay tracks 100 ms\nthen clears', C_ACTION)
    box(ax,    1.8, 1.8, 2.8, 0.5,  'Done (gear position\nupdated by CAN from rear)', C_TERM)

    # Manual shift down
    box(ax,    6.2, 4.7, 2.8, 0.5,  'SHIFT DOWN button pressed', C_MANUAL, bold=True)
    box(ax,    6.2, 3.8, 2.8, 0.55, 'Send CAN downshift\n100 ms pulse\ntargetGear = UNKNOWN', C_ACTION)
    box(ax,    6.2, 2.8, 2.8, 0.5,  'Relay tracks 100 ms\nthen clears', C_ACTION)
    box(ax,    6.2, 1.8, 2.8, 0.5,  'Done (gear position\nupdated by CAN from rear)', C_TERM)

    # Notes
    ax.text(4.0, 0.9, 'No clutch servo involvement in manual mode.', fontsize=7.5,
            ha='center', color='#404040', style='italic')
    ax.text(4.0, 0.55,'Neutral button has no function in manual mode.', fontsize=7.5,
            ha='center', color='#404040', style='italic')

    # arrows — toggle
    arrow(ax, 4.0, 10.05, 4.0, 9.65)
    arrow(ax, 4.0, 8.95,  4.0, 8.55, 'YES')
    ax.annotate('', xy=(4.0, 7.65), xytext=(4.0, 8.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(4.12, 8.2, 'NO', fontsize=7.5, color=C_ARROW)
    ax.annotate('', xy=(1.8, 6.55), xytext=(4.0, 6.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(2.5, 6.85, 'NO', fontsize=7.5, color=C_ARROW)
    ax.annotate('', xy=(6.2, 6.55), xytext=(4.0, 6.95),
                arrowprops=dict(arrowstyle='->', color=C_ARROW, lw=1.5))
    ax.text(5.2, 6.85, 'YES', fontsize=7.5, color=C_ARROW)

    # arrows — shift up
    arrow(ax, 1.8, 4.45, 1.8, 4.08)
    arrow(ax, 1.8, 3.52, 1.8, 3.05)
    arrow(ax, 1.8, 2.55, 1.8, 2.05)

    # arrows — shift down
    arrow(ax, 6.2, 4.45, 6.2, 4.08)
    arrow(ax, 6.2, 3.52, 6.2, 3.05)
    arrow(ax, 6.2, 2.55, 6.2, 2.05)

    return save_chart(fig, '_fc_manual.png')


# ─────────────────────────────────────────────────────────────────────────────
# Build Word doc
# ─────────────────────────────────────────────────────────────────────────────
charts = [
    ('UPSHIFT',   chart_upshift,   'Shift Up button pressed — ignition cut upshift'),
    ('DOWNSHIFT', chart_downshift, 'Shift Down button pressed — clutch servo downshift'),
    ('NEUTRAL',   chart_neutral,   'Neutral button pressed — single button, direction by gear'),
    ('MANUAL',    chart_manual,    'Manual mode — toggle and direct shift commands'),
]

doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin  = Cm(1.5)
section.bottom_margin = Cm(1.5)

for i, (name, fn, subtitle) in enumerate(charts):
    if i > 0:
        doc.add_page_break()
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.italic = True

    img_path = fn()
    doc.add_picture(img_path, width=Cm(16))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(img_path)

out = r'c:\Users\Adrian\Documents\PlatformIO\Projects\Rear node plus CanBUS\T89_Gearbox_Flowcharts.docx'
doc.save(out)
print('Saved: ' + out)
